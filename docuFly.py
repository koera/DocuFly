import pandas as pd
from docx import Document
from docx2pdf import convert
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import re
from string import Template
import json
import sys

def process_emails(config_path):
    # Load configuration
    with open(config_path, 'r') as config_file:
        config = json.load(config_file)

    excel_path = config['excel_path']
    word_template_path = config['word_template_path']
    output_dir = config['output_dir']
    email_credentials = config['email_credentials']
    email_subject = config['email_subject']
    file_name_pattern = config['file_name_pattern']
    email_body = config['email_body']
    email_cc = config['email_cc']
    
    # Load Excel file
    data = pd.read_excel(excel_path)
    
    # Gmail credentials
    gmail_user = email_credentials['username']
    gmail_password = email_credentials['password']
    
    # Extract placeholders from the Word template
    doc = Document(word_template_path)
    placeholders = set()
    for paragraph in doc.paragraphs:
        matches = re.findall(r'\$\{(.*?)\}', paragraph.text)
        placeholders.update(matches)
    
    # Ensure placeholders are valid column names in the Excel file
    valid_placeholders = {placeholder: placeholder for placeholder in placeholders if placeholder in data.columns}
    
    # Find the email column dynamically
    email_column = None
    for col in data.columns:
        if col.strip().lower() == 'email':
            email_column = col
            break

    if not email_column:
        print("Error: The 'email' column is required in the Excel file for sending emails.")
        sys.exit(1)

    valid_placeholders['email'] = email_column

    # Loop through each row in the Excel sheet
    for index, row in data.iterrows():
        # Dynamically map template variables to Excel data
        template_replacements = {placeholder: row[excel_field] for placeholder, excel_field in valid_placeholders.items()}
        recipient_email = row[valid_placeholders['email']]

        # Open the Word template
        doc = Document(word_template_path)

        # Replace placeholders with actual data
        for paragraph in doc.paragraphs:
            for placeholder, value in template_replacements.items():
                if f'${{{placeholder}}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace(f'${{{placeholder}}}', str(value))

        # Generate the file name dynamically
        file_name = Template(file_name_pattern).substitute(template_replacements)
        output_path = os.path.join(output_dir, file_name.replace('.pdf', '.docx'))

        # Save the document as Word
        doc.save(output_path)

        # Convert the Word document to PDF
        pdf_path = output_path.replace('.docx', '.pdf')
        convert(output_path)
        os.remove(output_path)

        # Prepare email
        msg = MIMEMultipart()
        msg['From'] = gmail_user
        msg['To'] = recipient_email
        msg['Subject'] = email_subject
        msg['cc'] = email_cc
        msg.attach(MIMEText(email_body, 'plain'))
        
        # Attach PDF file
        with open(pdf_path, "rb") as attachment:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename= {os.path.basename(pdf_path)}')
            msg.attach(part)
        
        # Send email
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(gmail_user, gmail_password)
        text = msg.as_string()
        server.sendmail(gmail_user, recipient_email, text)
        server.quit()
        print(f"Email sent to {recipient_email}")

    print("Documents created and emails sent successfully!")

if __name__ == '__main__':
    if len(sys.argv) != 2:
        print("Usage: python script.py <config_path>")
        sys.exit(1)
    
    config_path = sys.argv[1]  # Accept config file path as a command-line argument
    process_emails(config_path)
